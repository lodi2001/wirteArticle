import streamlit as st
from docx import Document
from openai import OpenAI
import pandas as pd
import matplotlib.pyplot as plt
from bertopic import BERTopic
from bertopic.representation import KeyBERTInspired
from functions import perform_topic_modelling,get_gpt_content
from openai import OpenAI
import io
from sklearn.decomposition import PCA
from sklearn.cluster import KMeans
from sklearn.feature_extraction.text import CountVectorizer

client =OpenAI(api_key="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")
# Initialize main app
st.set_page_config(page_title='Document Editor App', layout = 'wide')

# Create sidebar
st.sidebar.title("Navigation")
selection = st.sidebar.radio("Go to", ["Home", "Upload Article", "Write Literature Review", "Structure Paper"])

document = Document()
# Convert panda.Series into Array 

if selection == "Home":
    st.title("Home")
    # Home page content goes here

elif selection == "Upload Article":
    st.title("Upload Article")
    uploaded_file = st.file_uploader("Choose a file", type=['csv'])
    if uploaded_file is not None:
        # Assuming the uploaded file is a CSV, you can read it into a DataFrame
        data = pd.read_csv(uploaded_file)
        st.write(data)
        # Ensure you have a column named 'Abstract' in your CSV for BERTopic
        representation_model = KeyBERTInspired()
        model = BERTopic(representation_model=representation_model)
        if st.button('Generate Topic Clusters'):
            topics, probs = model.fit_transform(data['Abstract'])
            st.success("Topic clustering completed!")
            st.write(model.get_topic_info())
        if st.button('Generate topic by ChatGpt'):
            model=perform_topic_modelling(data['Abstract']).fit(data['Abstract'])
            st.write(model.get_topic_info())

elif selection == "Write Literature Review":
        
        st.title("Write Literature Review")
def convert_topic(documnets):
    
    print(documnets)
    documents_list = [item[0] for item in documnets]
    documents_list = documents_list[1:]
    single_text = ",".join(documents_list)
    return single_text
def convert_context(context):
    context_list=[item[0] for item in context]
    context_list = context_list[1:]
    return context_list
if st.button("Convert to Word Document"):
        combined_text = data['Title'] + " " + data['Abstract']

        dim_model = PCA(n_components=20)
        cluster_model = KMeans(n_clusters=30)
        vectorizer_model = CountVectorizer(ngram_range=(1, 2), stop_words="english")

        model = BERTopic(vectorizer_model=vectorizer_model,umap_model=dim_model, embedding_model="allenai-specter", 
                       hdbscan_model=cluster_model, calculate_probabilities= True,top_n_words=20)
        topics,probabilities = model.fit_transform(combined_text)

        # representation_model = KeyBERTInspired()
        # model = BERTopic(representation_model=representation_model)
        # combined_text = data['Title'] + " " + data['Abstract']
        # topics, probs = model.fit_transform(combined_text)
        st.success("Topic clustering completed!")
        #model.visualize_barchart(top_n_topics=49)
        fig = model.visualize_barchart()
        fig.write_html("C://dev//wirteArticle//file.html") # Use Streamlit to display the figure
        fig = model.visualize_barchart()

        # Use Streamlit to display the figure directly
        st.plotly_chart(fig)
        st.write(model.get_topic_info())
        documnets=model.get_topic_info()['Name']
        single_text= convert_topic(documnets)
        context=model.get_topic_info()['Representative_Docs']
        context_list=convert_context(context)
        text_input = get_gpt_content(single_text,context_list)
        print(text_input)
        document.add_paragraph(text_input)
        #document.add_paragraph(st.plotly_chart(fig))
        document.save("Literature_Review.docx")
        st.success("The text has been successfully saved to a Word document!")    

elif selection == "Structure Paper":
    st.title("Structure Paper")
    # Structure Paper page content goes here
